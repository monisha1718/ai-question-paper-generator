"""
Question paper generation module.

Builds a structured question paper (Parts A / B / C) using the RAG pipeline,
and exports the result to PDF or DOCX.
"""

import io
import logging
import textwrap
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from .rag_pipeline import retrieve_relevant_chunks, generate_with_context

logger = logging.getLogger(__name__)


def _build_prompt(
    subject: str,
    unit_or_topic: str,
    exam_type: str,
    marks_distribution: Optional[str] = None,
    regenerate: bool = False,
) -> str:
    """Construct the LLM prompt for generating a question paper."""
    regen_note = (
        "IMPORTANT: Generate completely NEW and DIFFERENT questions from any "
        "previously generated set. Vary the phrasing, focus areas, and difficulty.\n\n"
        if regenerate
        else ""
    )

    custom_marks = ""
    if marks_distribution:
        custom_marks = f"Additional marks guidance from the faculty: {marks_distribution}\n\n"

    prompt = textwrap.dedent(f"""\
        Generate an academic question paper with the following specifications:

        Subject     : {subject}
        Unit / Topic: {unit_or_topic}
        Exam Type   : {exam_type}

        {custom_marks}{regen_note}FORMAT (follow this EXACTLY):

        PART A - Short Answer Questions (2 marks each)
        Generate exactly 10 questions. Number them 1-10.

        PART B - Medium Answer Questions (5 marks each)
        Generate exactly 5 questions. Number them 11-15.

        PART C - Long Answer Questions (10 marks each)
        Generate exactly 3 questions. Number them 16-18.

        RULES:
        - All questions MUST be derived from the provided study material.
        - Use formal academic language.
        - Do NOT repeat or paraphrase the same question.
        - Cover as many different topics / concepts from the material as possible.
        - For PART C, include questions that require analysis, comparison, or detailed explanation.
        - Do not include answers.
    """)
    return prompt


def generate_question_paper(
    subject: str,
    unit_or_topic: str,
    exam_type: str,
    marks_distribution: Optional[str] = None,
    regenerate: bool = False,
) -> str:
    """
    Generate a question paper string using RAG.

    Returns the raw text of the question paper.
    """
    # Retrieve relevant chunks from the vector store
    query = f"{subject} {unit_or_topic} {exam_type} questions"
    chunks = retrieve_relevant_chunks(subject, query, k=12)

    if not chunks:
        return (
            "No study material found for this subject. "
            "Please upload a PDF before generating questions."
        )

    prompt = _build_prompt(subject, unit_or_topic, exam_type, marks_distribution, regenerate)
    paper = generate_with_context(chunks, prompt)
    return paper


# ── Export helpers ──────────────────────────────────────────────────────

def _header_block(subject: str, unit_or_topic: str, exam_type: str) -> str:
    return (
        f"Question Paper\n"
        f"Subject: {subject}\n"
        f"Unit / Topic: {unit_or_topic}\n"
        f"Exam Type: {exam_type}\n"
        f"{'=' * 60}\n\n"
    )


def export_to_pdf(
    paper_text: str,
    subject: str,
    unit_or_topic: str,
    exam_type: str,
) -> bytes:
    """Render the question paper text into a PDF and return the bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=0.75 * inch, bottomMargin=0.75 * inch,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "QPTitle", parent=styles["Title"], fontSize=16, spaceAfter=6,
        alignment=1,  # center
    )
    subtitle_style = ParagraphStyle(
        "QPSub", parent=styles["Normal"], fontSize=11, spaceAfter=4,
        alignment=1,
    )
    body_style = ParagraphStyle(
        "QPBody", parent=styles["Normal"], fontSize=11, leading=15,
        spaceAfter=4,
    )

    story = []
    story.append(Paragraph("Question Paper", title_style))
    story.append(Paragraph(f"Subject: {subject}", subtitle_style))
    story.append(Paragraph(f"Unit / Topic: {unit_or_topic}", subtitle_style))
    story.append(Paragraph(f"Exam Type: {exam_type}", subtitle_style))
    story.append(Spacer(1, 16))

    for line in paper_text.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 8))
            continue
        # Escape XML-special characters for ReportLab
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if line.upper().startswith("PART"):
            story.append(Paragraph(f"<b>{safe}</b>", body_style))
        else:
            story.append(Paragraph(safe, body_style))

    doc.build(story)
    return buf.getvalue()


def export_to_docx(
    paper_text: str,
    subject: str,
    unit_or_topic: str,
    exam_type: str,
) -> bytes:
    """Render the question paper text into a DOCX and return the bytes."""
    doc = DocxDocument()

    # Title
    title = doc.add_heading("Question Paper", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Subject: {subject}\n").bold = True
    meta.add_run(f"Unit / Topic: {unit_or_topic}\n")
    meta.add_run(f"Exam Type: {exam_type}\n")
    doc.add_paragraph("─" * 60)

    # Body
    for line in paper_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.upper().startswith("PART"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
